"""Load and manage the base term sheet template document."""
from __future__ import annotations

import os
import subprocess
from pathlib import Path
from typing import Optional

BASE_TEMPLATE_PATH = Path("/Users/pranjalsingh/Desktop/MASTERBASE_Series  Template Termsheet.doc")


def load_base_template() -> Optional[str]:
    """Load the base term sheet template document."""
    if not BASE_TEMPLATE_PATH.exists():
        return None
    
    try:
        # Try to read as .docx first (newer format)
        if BASE_TEMPLATE_PATH.suffix == ".docx":
            from docx import Document
            doc = Document(BASE_TEMPLATE_PATH)
            # Extract text from all paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            return "\n".join(paragraphs)
        elif BASE_TEMPLATE_PATH.suffix == ".doc":
            # For .doc files, try textutil (macOS) or catdoc
            try:
                # Try textutil first (macOS built-in)
                result = subprocess.run(
                    ["textutil", "-convert", "txt", "-stdout", str(BASE_TEMPLATE_PATH)],
                    capture_output=True,
                    text=True,
                    timeout=10,
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            
            try:
                # Try catdoc if available
                result = subprocess.run(
                    ["catdoc", str(BASE_TEMPLATE_PATH)],
                    capture_output=True,
                    text=True,
                    timeout=10,
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            
            return None
    except Exception:
        return None


def get_base_template_text() -> str:
    """Get base template text, with fallback to default structure."""
    template_text = load_base_template()
    if template_text and len(template_text) > 500:  # Ensure we got substantial content
        return template_text
    
    # Fallback: return a more complete template structure matching the Word document
    return """This indicative term sheet ("Term Sheet") summarizes the principal terms, for negotiation purposes only, with respect to a potential investment by [Insert Name of Investor] in [name of company]. Nothing herein creates any legally binding obligation on the part of any party, nor shall any legally binding obligations arise unless and until the parties have executed definitive written agreements (the "Definitive Agreements") and obtained all requisite governmental, corporate, management and legal approvals.

Company: [name of company] (the "Company"), having its registered office at [insert registered office]

Investor: [Insert Name of the Investor] and/or its affiliates (the "Investor").

Founder(s): [Insert Name of Founder 1], [Insert Name of Founder 2], [Insert Name of Founder 3] (Individually, each a "Founder", collectively "Founders")

Investment Amount and Investor Securities: The Investor shall invest [USD/INR] [Insert Investment Amount] into the Company through a primary investment at a [pre-money valuation/post money valuation] of [USD/INR] [Insert Amount] through a mix of compulsorily convertible cumulative preference shares ("CCCPS"), and a nominal number of equity shares; such that the Investor will own atleast [Insert shareholdng percentage]% of the Company on a fully diluted post-money basis, post investment. If the total round size increases, the premoney valuation of the company would be adjusted accordingly to give the shareholding percentage above.

ESOP: An ESOP pool will be included in the [pre-money valuation/post money valuation], equal to [Insert percentage]% of the fully diluted share capital of the Company on a pre-money basis."""

